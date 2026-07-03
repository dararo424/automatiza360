"""Generate docs/manual-usuario.docx from scratch using python-docx."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

INDIGO = RGBColor(0x4F, 0x46, 0xE5)
DARK   = RGBColor(0x0F, 0x17, 0x2A)
MUTED  = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
GRAY_BG = RGBColor(0xF8, 0xFA, 0xFC)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text='', bold=False, italic=False, size=11,
                  color=None, space_before=0, space_after=6, align=None,
                  indent=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent:
        pf.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    return p


def add_run(para, text, bold=False, italic=False, size=11, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_heading1(doc, text):
    """Section header."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # Shade background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '4F46E5')
    pPr.append(shd)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = WHITE
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    pf = p.paragraph_format
    pf.left_indent = Cm(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = DARK
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E2E8F0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = INDIGO
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_numbered(doc, items):
    for i, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{i}. {text}")
        run.font.size = Pt(11)


def add_callout(doc, icon, text, bg='EFF6FF'):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # icon cell
    cell0 = table.cell(0, 0)
    cell0.width = Cm(1)
    set_cell_bg(cell0, bg)
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(icon)
    r0.font.size = Pt(14)
    # text cell
    cell1 = table.cell(0, 1)
    set_cell_bg(cell1, bg)
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run(text)
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK
    doc.add_paragraph()  # spacer


def add_simple_table(doc, headers, rows):
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_bg(cell, '4F46E5')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.cell(ri + 1, ci)
            if ri % 2 == 1:
                set_cell_bg(cell, 'F8FAFC')
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()


def add_screen_box(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, '1E293B')
    for line in lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        r.font.name = 'Courier New'
    # remove first empty paragraph
    first = cell.paragraphs[0]
    if not first.text:
        first._element.getparent().remove(first._element)
    doc.add_paragraph()


# ──────────────────────────────────────────────
# BUILD DOCUMENT
# ──────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Normal style defaults
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── COVER ──────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A360')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = INDIGO

add_paragraph(doc, 'Manual de Usuario', bold=True, size=24,
              color=DARK, space_after=4, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, 'Automatiza360 — Gestiona tu negocio desde WhatsApp',
              size=14, color=MUTED, space_after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, 'Versión 1.0  ·  Junio 2026  ·  automatiza360.com',
              size=10, color=MUTED, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ── TABLE OF CONTENTS ─────────────────────────
add_paragraph(doc, '📋 Contenido', bold=True, size=14, color=DARK,
              space_before=0, space_after=8)
toc_items = [
    '1. Primeros pasos — Registro y configuración',
    '2. Panel principal (Dashboard)',
    '3. Bot de WhatsApp para clientes',
    '4. Bot de administración (WhatsApp)',
    '5. Gestión de pedidos y órdenes',
    '6. Agenda y citas',
    '7. Tickets de reparación',
    '8. Productos e inventario',
    '9. Contactos y programa de lealtad',
    '10. Campañas de WhatsApp',
    '11. Gastos y caja',
    '12. Equipo y usuarios',
    '13. Integraciones (Instagram DMs)',
    '14. Configuración del negocio',
    '15. Planes y facturación',
    '16. Preguntas frecuentes',
]
for item in toc_items:
    add_bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 1 — 🚀 Primeros pasos — Registro y configuración')

add_heading2(doc, '1.1 Crear tu cuenta')
add_paragraph(doc, 'El proceso de registro toma menos de 3 minutos y no requiere tarjeta de crédito. Tienes 14 días de prueba gratuita.')
add_numbered(doc, [
    'Ingresa a automatiza360.com y haz clic en "Crear cuenta gratis".',
    'Escribe el nombre de tu negocio y tu nombre completo.',
    'Selecciona la industria de tu negocio (restaurante, clínica, tienda, etc.).',
    'Ingresa tu correo electrónico y crea una contraseña segura.',
    'Acepta los términos y haz clic en "Crear mi negocio".',
])
add_callout(doc, 'ℹ️', 'Recibirás un correo de bienvenida con los siguientes pasos para configurar tu bot de WhatsApp.')

add_heading2(doc, '1.2 Checklist de onboarding')
add_paragraph(doc, 'Al ingresar por primera vez verás una lista de tareas en tu panel principal. Completa cada paso para tener tu negocio listo:')
add_simple_table(doc,
    ['Paso', 'Descripción'],
    [
        ['📝 Perfil del negocio', 'Nombre, descripción, horarios y dirección'],
        ['📦 Agregar productos', 'Crea tu catálogo o menú inicial'],
        ['👥 Invitar equipo', 'Agrega a tus empleados o socios'],
        ['📱 Activar WhatsApp', 'Conecta tu número de WhatsApp Business'],
        ['🔗 Compartir enlace', 'Difunde tu página pública y código QR'],
    ]
)

add_heading2(doc, '1.3 Tu página pública y código QR')
add_paragraph(doc, 'Automatiza360 crea automáticamente una página pública para tu negocio con la URL:')
add_screen_box(doc, ['https://automatiza360.com/negocio/tu-negocio'])
add_paragraph(doc, 'Esta página muestra tu información, horarios y un botón directo a WhatsApp. Desde el panel puedes descargar tu código QR para imprimir y poner en tu negocio.')

# ══════════════════════════════════════════════
# SECTION 2
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 2 — 📊 Panel principal (Dashboard)')
add_paragraph(doc, 'El dashboard es la pantalla principal de Automatiza360. Muestra un resumen en tiempo real del estado de tu negocio.')

add_heading2(doc, '2.1 Métricas del día')
add_bullet(doc, '🛒 Órdenes / Citas — Total del día vs. ayer para ver si estás por encima o debajo.')
add_bullet(doc, '💰 Ingresos del mes — Suma de todas las ventas no canceladas del mes en curso.')
add_bullet(doc, '💬 Conversaciones — Clientes atendidos por el bot este mes vs. tu cuota del plan.')
add_bullet(doc, '⚠️ Stock bajo — Productos por debajo del stock mínimo configurado.')

add_heading2(doc, '2.2 Gráfica de tendencias')
add_paragraph(doc, 'La gráfica muestra la evolución de ventas, citas e ingresos de los últimos 30 días. Úsala para identificar los días pico de tu negocio.')

add_heading2(doc, '2.3 Widget de ROI')
add_paragraph(doc, 'Este widget calcula cuánto tiempo y dinero te ha ahorrado el bot este mes:')
add_bullet(doc, 'Mensajes automatizados: total de respuestas enviadas por el bot')
add_bullet(doc, 'Horas ahorradas: estimado en base a 3 minutos por mensaje')
add_bullet(doc, 'Ahorro estimado: en pesos colombianos')

add_heading2(doc, '2.4 Últimas actividades')
add_paragraph(doc, 'Ver las últimas 5 órdenes y próximas citas sin tener que navegar a cada módulo.')

# ══════════════════════════════════════════════
# SECTION 3
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 3 — 🤖 Bot de WhatsApp para clientes')
add_paragraph(doc, 'El bot de WhatsApp es tu empleado virtual disponible 24 horas, 7 días a la semana. Atiende a tus clientes, toma pedidos, agenda citas y responde preguntas — todo de forma automática.')

add_heading2(doc, '3.1 ¿Qué puede hacer el bot?')
add_simple_table(doc,
    ['Industria', 'Funciones principales'],
    [
        ['🍽️ Restaurante / Panadería', 'Mostrar menú del día, tomar pedidos completos, confirmar pago'],
        ['🔧 Tienda Tech / Taller', 'Registrar equipos para reparación, consultar estado, cotizaciones'],
        ['🏥 Clínica / Consultorio', 'Ver disponibilidad, agendar citas, enviar recordatorios'],
        ['💅 Salón / Spa / Veterinaria', 'Reservar por profesional y servicio, modificar citas'],
        ['👗 Tienda de ropa', 'Consultar tallas, catálogo de productos, hacer pedidos'],
        ['💊 Farmacia', 'Consultar disponibilidad de medicamentos, precios'],
    ]
)

add_heading2(doc, '3.2 Funciones comunes (todos los negocios)')
add_bullet(doc, 'Notas de voz: el bot transcribe y entiende los audios de los clientes')
add_bullet(doc, 'Cupones de descuento: valida y aplica códigos automáticamente')
add_bullet(doc, 'Puntos de lealtad: informa al cliente sus puntos acumulados')
add_bullet(doc, 'Escalación humana: si el cliente pide un humano, te avisa en tiempo real')
add_bullet(doc, 'Reinicio: el cliente escribe "menu" o "reiniciar" para empezar de nuevo')

add_heading2(doc, '3.3 Ejemplo de conversación (Restaurante)')
add_screen_box(doc, [
    '👤 Cliente: hola, qué hay hoy?',
    '🤖 Bot: ¡Hola! Hoy tenemos:',
    '   • Bandeja paisa — $18.000',
    '   • Sancocho de gallina — $15.000',
    '   • Ajiaco — $16.000',
    '   ¿Qué te gustaría pedir? 😊',
    '',
    '👤 Cliente: una bandeja paisa por favor, pago con efectivo',
    '🤖 Bot: ¡Perfecto! Resumen de tu pedido:',
    '   • Bandeja paisa x1 — $18.000',
    '   Pago: Efectivo. ¿Confirmas el pedido?',
    '',
    '👤 Cliente: sí',
    '🤖 Bot: ✅ Pedido #47 confirmado. Te avisamos cuando esté listo.',
    '   Powered by Automatiza360 🤖',
])

add_heading2(doc, '3.4 Opt-out (desuscripción)')
add_paragraph(doc, 'Si un cliente no desea recibir más mensajes, puede escribir STOP en cualquier momento. El sistema lo registra automáticamente y no le enviarás más campañas a ese número.')
add_callout(doc, 'ℹ️', 'El bot responde en el mismo idioma y tono configurado para tu negocio. Puedes ajustar el nombre del bot y su tono (Formal / Amigable / Costeño) desde Configuración.')

# ══════════════════════════════════════════════
# SECTION 4
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 4 — 🧑‍💼 Bot de administración (WhatsApp)')
add_paragraph(doc, 'El bot de administración te permite gestionar tu negocio desde el mismo número de WhatsApp, sin abrir ninguna app. Escríbele como si fuera un asistente personal.')
add_callout(doc, '💡', 'El sistema detecta automáticamente si quien escribe es dueño, admin o empleado, y activa el modo de administración. Tus clientes nunca ven estas conversaciones.', bg='F0FDF4')

add_heading2(doc, '4.1 Comandos disponibles')
add_simple_table(doc,
    ['Lo que escribes', 'Lo que hace el bot'],
    [
        ['"resumen del día"', 'Muestra ingresos, órdenes, citas y gastos de hoy'],
        ['"órdenes pendientes"', 'Lista las órdenes que faltan por preparar'],
        ['"citas de mañana"', 'Muestra el horario completo del día siguiente'],
        ['"gasté $45.000 en ingredientes"', 'Registra el gasto automáticamente'],
        ['"hoy hay bandeja paisa $18.000"', 'Actualiza el menú del día (restaurantes)'],
        ['"estado del ticket de Juan García"', 'Busca y muestra el ticket de reparación'],
        ['"cancelar citas de mañana desde las 2pm"', 'Cancela en rango y notifica a los pacientes'],
        ['"últimas reseñas"', 'Muestra valoraciones recientes de clientes'],
        ['"crear cupón 20% descuento"', 'Crea un código de descuento nuevo'],
    ]
)

add_heading2(doc, '4.2 Enviar reportes por WhatsApp o email')
add_bullet(doc, 'Gráfica por WhatsApp: imagen con las ventas de la semana')
add_bullet(doc, 'Reporte ejecutivo por email: PDF con métricas del mes')

add_heading2(doc, '4.3 Cambiar estado de órdenes y citas')
add_screen_box(doc, [
    '👤 Dueño: orden 47 está lista',
    '🤖 Admin Bot: ✅ Orden #47 marcada como LISTA. El cliente fue notificado por WhatsApp automáticamente.',
])

# ══════════════════════════════════════════════
# SECTION 5
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 5 — 🛒 Gestión de pedidos y órdenes')

add_heading2(doc, '5.1 Ver pedidos')
add_paragraph(doc, 'Ve a Órdenes en el menú lateral. Verás todos los pedidos ordenados por fecha, con filtros por estado:')
add_bullet(doc, 'PENDIENTE — Pedido recibido, sin confirmar')
add_bullet(doc, 'CONFIRMADO — Confirmado, en preparación')
add_bullet(doc, 'PREPARANDO — En cocina o proceso')
add_bullet(doc, 'LISTO — Listo para entrega/recogida')
add_bullet(doc, 'ENTREGADO — Completado')

add_heading2(doc, '5.2 Cambiar estado')
add_paragraph(doc, 'Haz clic en cualquier orden y selecciona el nuevo estado. Al pasar a LISTO, el cliente recibe automáticamente un WhatsApp de notificación.')

add_heading2(doc, '5.3 Detalle de orden')
add_paragraph(doc, 'Cada orden muestra: lista de ítems, cantidades, total, método de pago, teléfono del cliente y historial de cambios de estado.')
add_callout(doc, '⚠️', 'Las órdenes canceladas no se eliminan — quedan en el historial para efectos contables. Puedes filtrarlas para no verlas en la vista principal.', bg='FFFBEB')

# ══════════════════════════════════════════════
# SECTION 6
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 6 — 📅 Agenda y citas')

add_heading2(doc, '6.1 Ver agenda')
add_paragraph(doc, 'Ve a Agenda en el menú. Puedes ver las citas en formato lista, día o semana, y filtrar por profesional o servicio.')

add_heading2(doc, '6.2 Crear cita manualmente')
add_numbered(doc, [
    'Haz clic en el botón "Nueva cita" (esquina superior derecha).',
    'Selecciona el servicio y el profesional.',
    'Elige la fecha y hora disponible.',
    'Ingresa el nombre y teléfono del cliente.',
    'Guarda — el cliente recibe confirmación por WhatsApp.',
])

add_heading2(doc, '6.3 Recordatorios automáticos')
add_paragraph(doc, 'El sistema envía automáticamente un recordatorio por WhatsApp 24 horas antes de cada cita. No necesitas hacer nada.')

add_heading2(doc, '6.4 Estados de citas')
add_simple_table(doc,
    ['Estado', 'Significado'],
    [
        ['AGENDADA', 'Cita creada, pendiente de confirmación'],
        ['CONFIRMADA', 'Cliente confirmó asistencia'],
        ['COMPLETADA', 'Servicio realizado'],
        ['NO ASISTIÓ', 'Cliente no se presentó (no-show)'],
        ['CANCELADA', 'Cancelada por el negocio o el cliente'],
    ]
)

add_heading2(doc, '6.5 Cancelación masiva')
add_screen_box(doc, [
    '👤 "cancela las citas del viernes desde las 2pm"',
    '🤖 "¿Cancelo 4 citas del viernes 16 de mayo desde las 2:00 PM?',
    '   (Clientes: María, Pedro, Luisa, Carlos)"',
    '👤 "sí"',
    '🤖 "✅ Cancelé 4 citas y notifiqué a los pacientes por WhatsApp."',
])

# ══════════════════════════════════════════════
# SECTION 7
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 7 — 🔧 Tickets de reparación')
add_paragraph(doc, 'Módulo diseñado para talleres y tiendas tech. Gestiona el flujo completo desde que el cliente deja el equipo hasta que lo recoge.')

add_heading2(doc, '7.1 Crear un ticket')
add_numbered(doc, [
    'Ve a Tickets y haz clic en "Nuevo ticket".',
    'Ingresa datos del cliente (nombre, teléfono) y del equipo (dispositivo, falla reportada).',
    'Opcionalmente adjunta fotos del equipo.',
    'El ticket se crea con estado RECIBIDO y se le asigna un número único.',
])

add_heading2(doc, '7.2 Flujo de estados')
add_screen_box(doc, ['RECIBIDO → DIAGNÓSTICO → ESPERANDO PIEZAS → REPARANDO → LISTO → ENTREGADO'])

add_heading2(doc, '7.3 Garantías')
add_paragraph(doc, 'Al marcar un ticket como ENTREGADO, puedes definir el período de garantía. El sistema guarda las garantías vigentes y el bot puede consultarlas cuando el cliente regrese.')

add_heading2(doc, '7.4 Crear tickets desde WhatsApp (Admin Bot)')
add_screen_box(doc, [
    '👤 "nuevo ticket: Juan García, cel 3001234567, Samsung S22, no enciende"',
    '🤖 "✅ Ticket #89 creado para Juan García — Samsung S22.',
    '   Si tienes foto del equipo, mándala y la adjunto."',
])

# ══════════════════════════════════════════════
# SECTION 8
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 8 — 📦 Productos e inventario')

add_heading2(doc, '8.1 Agregar un producto')
add_numbered(doc, [
    'Ve a Productos o Inventario.',
    'Haz clic en "Nuevo producto".',
    'Completa: nombre, precio, stock actual, stock mínimo y categoría.',
    'Sube una foto (opcional).',
    'Guarda — el producto ya está disponible para el bot.',
])

add_heading2(doc, '8.2 Alertas de stock bajo')
add_paragraph(doc, 'Cuando el stock de un producto cae por debajo del stock mínimo configurado, aparece una alerta en el dashboard. El bot de admin también puede listar los productos con stock bajo al escribir "stock bajo".')

add_heading2(doc, '8.3 Importar inventario masivo')
add_paragraph(doc, 'Puedes importar tu catálogo completo desde:')
add_bullet(doc, 'Archivo Excel (.xlsx)')
add_bullet(doc, 'Archivo PDF con lista de productos')
add_bullet(doc, 'Foto de una lista escrita o impresa')
add_paragraph(doc, 'El sistema usa inteligencia artificial para extraer los productos automáticamente.')

add_heading2(doc, '8.4 Menú del día (Restaurantes)')
add_paragraph(doc, 'Ve a Menú del día para publicar los platos disponibles hoy con sus precios. El bot los mostrará automáticamente a los clientes que pregunten.')

# ══════════════════════════════════════════════
# SECTION 9
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 9 — 👥 Contactos y programa de lealtad')

add_heading2(doc, '9.1 Base de contactos')
add_paragraph(doc, 'Todos los clientes que han interactuado con tu bot se guardan automáticamente como contactos. Puedes ver su historial de compras, citas y conversaciones.')

add_heading2(doc, '9.2 Programa de puntos')
add_paragraph(doc, 'Cada compra acumula puntos automáticamente (1 punto por cada $10.000 COP):')
add_simple_table(doc,
    ['Acción', 'Puntos'],
    [
        ['Compra de $10.000', '+1 punto'],
        ['Compra de $100.000', '+10 puntos'],
        ['Canje de puntos', 'Se descuentan del total'],
    ]
)
add_paragraph(doc, 'El bot informa al cliente sus puntos disponibles y les permite canjearlos en su próxima compra.')

add_heading2(doc, '9.3 Segmentación por etiquetas')
add_paragraph(doc, 'Puedes asignar etiquetas a tus contactos (ej: VIP, frecuente, nuevo) y usar esas etiquetas para enviar campañas segmentadas.')

add_heading2(doc, '9.4 Fecha de cumpleaños')
add_paragraph(doc, 'Al agregar la fecha de cumpleaños de un contacto, el sistema envía automáticamente un mensaje de felicitación el día de su cumpleaños.')

# ══════════════════════════════════════════════
# SECTION 10
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 10 — 📢 Campañas de WhatsApp')
add_callout(doc, 'ℹ️', 'Las campañas de WhatsApp están disponibles en los planes Pro y Business.')

add_heading2(doc, '10.1 Crear una campaña')
add_numbered(doc, [
    'Ve a Campañas y haz clic en "Nueva campaña".',
    'Escribe el mensaje. Usa {nombre} para personalizar con el nombre del cliente.',
    'Define los filtros de segmentación (opcional).',
    'Haz clic en "Ver destinatarios" para confirmar cuántos recibirán el mensaje.',
    'Haz clic en "Enviar campaña".',
])

add_heading2(doc, '10.2 Filtros de segmentación')
add_simple_table(doc,
    ['Filtro', 'Uso'],
    [
        ['Por etiquetas', 'Solo contactos con etiqueta específica'],
        ['Sin compras en N días', 'Reactivar clientes inactivos'],
        ['Compras recientes', 'Premiar a clientes frecuentes'],
        ['Puntos mínimos', 'Campañas exclusivas para clientes VIP'],
    ]
)

add_heading2(doc, '10.3 Ejemplo de mensaje')
add_screen_box(doc, [
    '¡Hola {nombre}! 👋 Este fin de semana tenemos 20% de descuento en todos nuestros servicios.',
    'Usa el código FINDE20 al agendar tu cita.',
    '¡Te esperamos! 😊',
    '',
    'Responde STOP para no recibir más mensajes.',
])

add_heading2(doc, '10.4 Cupones de descuento')
add_paragraph(doc, 'Ve a Cupones para crear códigos de descuento. Puedes definir:')
add_bullet(doc, 'Tipo: porcentaje (%) o monto fijo ($)')
add_bullet(doc, 'Fecha de vencimiento')
add_bullet(doc, 'Código personalizado')

# ══════════════════════════════════════════════
# SECTION 11
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 11 — 💵 Gastos y caja')

add_heading2(doc, '11.1 Registrar gastos')
add_paragraph(doc, 'Ve a Gastos para registrar los costos del negocio. Puedes categorizarlos (insumos, servicios, nómina, etc.) y ver el comparativo ingresos vs. gastos del mes.')

add_heading2(doc, '11.2 Registrar desde WhatsApp')
add_screen_box(doc, [
    '👤 "gasté $80.000 en ingredientes para la cocina"',
    '🤖 "✅ Gasto de $80.000 registrado en categoría Insumos."',
])

add_heading2(doc, '11.3 Módulo de caja')
add_paragraph(doc, 'El módulo de Caja muestra el flujo de dinero del día: cuánto entró (ventas) y cuánto salió (gastos), con el saldo neto.')

# ══════════════════════════════════════════════
# SECTION 12
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 12 — 👨‍👩‍👧 Equipo y usuarios')

add_heading2(doc, '12.1 Roles disponibles')
add_simple_table(doc,
    ['Rol', 'Permisos'],
    [
        ['🔑 Dueño (Owner)', 'Acceso total, configuración, facturación'],
        ['⚙️ Admin', 'Gestión completa excepto facturación'],
        ['👤 Staff', 'Operación diaria, sin configuración'],
    ]
)

add_heading2(doc, '12.2 Invitar un miembro')
add_numbered(doc, [
    'Ve a Equipo y haz clic en "Invitar miembro".',
    'Ingresa el correo electrónico y selecciona el rol.',
    'La persona recibirá un correo de invitación para crear su cuenta.',
])

add_heading2(doc, '12.3 Límites por plan')
add_simple_table(doc,
    ['Plan', 'Máximo de miembros'],
    [
        ['Starter', '3 miembros'],
        ['Pro', '10 miembros'],
        ['Business', 'Ilimitados'],
    ]
)

add_heading2(doc, '12.4 Turnos del personal')
add_paragraph(doc, 'Ve a Turnos para gestionar los horarios de tu equipo. El bot de admin puede consultarlos escribiendo "turnos de hoy".')

# ══════════════════════════════════════════════
# SECTION 13
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 13 — 📷 Integraciones — Instagram DMs')
add_paragraph(doc, 'Conecta tu cuenta de Instagram Business para que el bot responda automáticamente los mensajes directos (DMs) de tus seguidores — exactamente igual que WhatsApp.')

add_heading2(doc, '13.1 Requisitos previos')
add_bullet(doc, 'Cuenta de Instagram Business (no cuenta personal)')
add_bullet(doc, 'La cuenta de Instagram debe estar vinculada a una Página de Facebook')
add_callout(doc, 'ℹ️', 'Si tienes una cuenta personal de Instagram, puedes convertirla a Business gratis desde la configuración de Instagram: Configuración → Cuenta → Cambiar a cuenta profesional.')

add_heading2(doc, '13.2 Conectar Instagram')
add_numbered(doc, [
    'Ve a Configuración en el menú lateral.',
    'Busca la sección "Instagram DMs".',
    'Haz clic en el botón "Conectar Instagram →".',
    'Se abrirá una ventana de Meta — inicia sesión con tu cuenta de Facebook.',
    'Selecciona los permisos y haz clic en "Continuar".',
    'Regresarás al panel con la confirmación "✅ Instagram conectado exitosamente".',
])

add_heading2(doc, '13.3 ¿Qué pasa después?')
add_paragraph(doc, 'Una vez conectado, el bot responderá todos los DMs nuevos que lleguen a tu Instagram con las mismas capacidades que el bot de WhatsApp. Las conversaciones también quedan registradas en el panel.')

add_heading2(doc, '13.4 Desconectar Instagram')
add_paragraph(doc, 'En cualquier momento puedes desconectar tu Instagram desde la misma sección de Configuración. Haz clic en "Desconectar Instagram".')

# ══════════════════════════════════════════════
# SECTION 14
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 14 — ⚙️ Configuración del negocio')

add_heading2(doc, '14.1 Información básica')
add_paragraph(doc, 'Ve a Configuración para actualizar:')
add_bullet(doc, 'Nombre del negocio y descripción')
add_bullet(doc, 'Horarios de atención')
add_bullet(doc, 'Dirección y ciudad')
add_bullet(doc, 'Logo (foto de perfil)')

add_heading2(doc, '14.2 Personalizar el bot')
add_simple_table(doc,
    ['Opción', 'Descripción'],
    [
        ['Nombre del bot', 'Cómo se presenta el asistente (ej: "Soy Sofía, tu asistente...")'],
        ['Tono Formal', 'Profesional y respetuoso, ideal para clínicas y servicios'],
        ['Tono Amigable', 'Cercano y cálido, recomendado para la mayoría'],
        ['Tono Costeño', 'Descomplicado y alegre, ideal para negocios del Caribe'],
    ]
)

add_heading2(doc, '14.3 Sucursales')
add_paragraph(doc, 'Si tienes varias ubicaciones, ve a Sucursales para registrarlas. Cada sucursal puede tener su propio horario y dirección.')

add_heading2(doc, '14.4 Recuperar contraseña')
add_paragraph(doc, 'Si olvidaste tu contraseña, haz clic en "¿Olvidaste tu contraseña?" en la pantalla de inicio de sesión. Recibirás un correo con un enlace válido por 1 hora.')

# ══════════════════════════════════════════════
# SECTION 15
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 15 — 💳 Planes y facturación')

add_heading2(doc, '15.1 Planes disponibles')
add_simple_table(doc,
    ['Plan', 'Precio mensual', 'Conversaciones', 'Equipo', 'Campañas'],
    [
        ['Starter', '$79.000 COP', '500/mes', '3 personas', 'No'],
        ['Pro', '$242.000 COP', '2.000/mes', '10 personas', 'Sí'],
        ['Business', '$529.000 COP', 'Ilimitadas', 'Ilimitado', 'Sí'],
    ]
)

add_heading2(doc, '15.2 Actualizar tu plan')
add_numbered(doc, [
    'Ve a Mi Plan en el menú lateral.',
    'Selecciona el plan al que deseas pasar.',
    'Completa el pago con tu tarjeta (procesado por Wompi — 100% seguro).',
    'Tu plan se activa inmediatamente y recibes el recibo por correo.',
])

add_heading2(doc, '15.3 Período de prueba')
add_paragraph(doc, 'Todos los negocios nuevos tienen 14 días de prueba gratuita con acceso completo al plan Pro. No se requiere tarjeta de crédito. Al terminar el período de prueba, el plan pasa a Starter automáticamente o puedes actualizar antes de que expire.')

add_heading2(doc, '15.4 Reembolsos')
add_paragraph(doc, 'Si tienes problemas con un pago, contacta a soporte en soporteautomatiza360@rgytgroup.com dentro de los primeros 7 días del cobro.')

# ══════════════════════════════════════════════
# SECTION 16
# ══════════════════════════════════════════════
add_heading1(doc, 'Sección 16 — ❓ Preguntas frecuentes')

faq = [
    ('¿Necesito un número de WhatsApp nuevo?',
     'El equipo de Automatiza360 te asigna un número de WhatsApp Business. Si ya tienes uno aprobado por Meta, también se puede usar.'),
    ('¿El bot funciona las 24 horas?',
     'Sí. El bot nunca descansa. Atiende a tus clientes en cualquier horario, incluso domingos y festivos.'),
    ('¿Los clientes saben que están hablando con un bot?',
     'Los mensajes incluyen "Powered by Automatiza360 🤖" al final. Puedes ajustar el nombre del bot para que se sienta más personal.'),
    ('¿Puedo seguir respondiendo manualmente?',
     'Sí. En la sección Conversaciones puedes ver todos los chats y responder directamente si lo prefieres. Si un cliente pide hablar con un humano, recibirás una alerta.'),
    ('¿Qué pasa si se me acaba la cuota de conversaciones?',
     'El bot le informa al cliente que el negocio ha alcanzado su límite del mes y lo invita a contactar en el próximo mes. Puedes actualizar tu plan en cualquier momento para evitarlo.'),
    ('¿Mis datos están seguros?',
     'Sí. Cada negocio tiene sus datos completamente aislados. Usamos cifrado, contraseñas hasheadas y conexiones seguras (HTTPS). Cumplimos con la Ley 1581 de Protección de Datos de Colombia.'),
    ('¿Cómo contacto a soporte?',
     'Escríbenos a soporteautomatiza360@rgytgroup.com o abre un ticket desde el panel en la sección Soporte. Tiempo de respuesta: máximo 24 horas hábiles.'),
]
for q, a in faq:
    add_heading3(doc, q)
    add_paragraph(doc, a)

# ── FOOTER ─────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Automatiza360')
r.bold = True
r.font.color.rgb = INDIGO
r.font.size = Pt(14)

add_paragraph(doc, 'Manual de Usuario v1.0  ·  Junio 2026',
              size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '¿Tienes dudas? Escríbenos a soporteautomatiza360@rgytgroup.com',
              size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, '© 2026 Automatiza360. Todos los derechos reservados.',
              size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── SAVE ──────────────────────────────────────
import os
os.makedirs('/home/user/automatiza360/docs', exist_ok=True)
doc.save('/home/user/automatiza360/docs/manual-usuario.docx')
print('✅ manual-usuario.docx created successfully')
